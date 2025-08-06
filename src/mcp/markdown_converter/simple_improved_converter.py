"""Simplified improved markdown to DOCX conversion using available libraries."""

from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX
from pathlib import Path
import markdown
import xml.etree.ElementTree as ET
import re
from .template_manager import TemplateManager
from .image_handler import ImageHandler
from .mermaid_handler import MermaidHandler
from .style_mapper import StyleMapper, MarkdownElement


class SimpleImprovedConverter:
    """Improved markdown to DOCX converter using standard markdown library."""
    
    def __init__(self, template_path: Optional[str] = None, template_dirs: Optional[list] = None, base_path: Optional[str] = None):
        self.template_path = template_path
        self.template_manager = TemplateManager(template_dirs)
        self.image_handler = ImageHandler(base_path)
        self.mermaid_handler = MermaidHandler()
        self.style_mapper = StyleMapper()
        
        # Configure markdown parser with available extensions
        self.md = markdown.Markdown(extensions=[
            'tables',           # GFM tables
            'fenced_code',      # Fenced code blocks
            'codehilite',       # Code highlighting
            'toc',              # Table of contents
            'nl2br',            # Newline to break
        ], extension_configs={
            'codehilite': {
                'use_pygments': False,
                'noclasses': True,
            }
        })
        
    def convert(self, markdown_content: str, metadata: Optional[Dict[str, Any]] = None, template_name: Optional[str] = None) -> Document:
        """Convert markdown content to Word document using proper parsing."""
        template_path = None
        if template_name:
            template_path = self.template_manager.get_template_path(template_name)
        elif self.template_path:
            template_path = self.template_path
            
        if template_path and self.template_manager.validate_template(template_path):
            doc = Document(template_path)
        else:
            doc = Document()
            
        if metadata:
            self._add_metadata(doc, metadata)
            
        # Pre-process for task lists and strikethrough
        processed_content = self._preprocess_markdown(markdown_content)
        
        # Process images and mermaid diagrams
        processed_content = self.image_handler.process_images(doc, processed_content)
        processed_content = self.mermaid_handler.process_mermaid(doc, processed_content)
        
        # Parse markdown to HTML
        html_content = self.md.convert(processed_content)
        
        # Convert HTML to Word document
        self._convert_html_to_docx(doc, html_content)
        
        return doc
    
    def _preprocess_markdown(self, content: str) -> str:
        """Pre-process markdown for features not supported by standard extensions."""
        # Remove YAML front matter
        content = re.sub(r'^---\s*\n.*?\n---\s*\n', '', content, flags=re.DOTALL)
        
        # Handle task lists
        content = re.sub(r'^(\s*)- \[x\]', r'\\1- ☑', content, flags=re.MULTILINE)
        content = re.sub(r'^(\s*)- \[ \]', r'\\1- ☐', content, flags=re.MULTILINE)
        
        # Handle strikethrough (convert to HTML)
        content = re.sub(r'~~([^~]+)~~', r'<del>\\1</del>', content)
        
        return content
    
    def _add_metadata(self, doc: Document, metadata: Dict[str, Any]) -> None:
        """Add document metadata."""
        core_props = doc.core_properties
        if 'title' in metadata:
            core_props.title = metadata['title']
        if 'author' in metadata:
            core_props.author = metadata['author']
        if 'subject' in metadata:
            core_props.subject = metadata['subject']
        if 'keywords' in metadata:
            core_props.keywords = metadata['keywords']
        if 'comments' in metadata:
            core_props.comments = metadata['comments']
    
    def _convert_html_to_docx(self, doc: Document, html_content: str) -> None:
        """Convert HTML to Word document elements."""
        try:
            # Clean up HTML for parsing
            html_content = html_content.replace('&nbsp;', ' ')
            html_content = re.sub(r'<br\s*/?>', '<br/>', html_content)
            html_content = re.sub(r'<hr\s*/?>', '<hr/>', html_content)
            
            # Handle HTML entities more carefully
            html_content = html_content.replace('&amp;', '&')
            html_content = html_content.replace('&lt;', '<')
            html_content = html_content.replace('&gt;', '>')
            
            # Parse HTML
            root = ET.fromstring(f'<root>{html_content}</root>')
            
            # Process all child elements
            for child in root:
                self._process_html_element(doc, child)
                
            # Handle any remaining text content
            if root.text and root.text.strip():
                doc.add_paragraph(root.text.strip())
                
        except ET.ParseError as e:
            print(f"HTML parsing error: {e}")
            print(f"Problematic HTML: {html_content[:200]}...")
            # Better fallback - try to extract readable content
            self._fallback_text_conversion(doc, html_content)
    
    def _process_html_element(self, doc: Document, element: ET.Element, parent_para=None) -> None:
        """Process HTML elements recursively."""
        tag = element.tag.lower()
        
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(tag[1])
            text = self._get_element_text(element)
            doc.add_heading(text, level=level)
            
        elif tag == 'p':
            para = doc.add_paragraph()
            # Apply paragraph style
            word_style = self.style_mapper.get_style_for_html_tag(tag)
            if word_style:
                self.style_mapper.apply_style_to_paragraph(para, word_style)
            self._process_inline_elements(para, element)
            
        elif tag == 'pre':
            code_element = element.find('code')
            if code_element is not None:
                self._add_code_block(doc, code_element)
            else:
                self._add_code_block(doc, element)
                
        elif tag == 'table':
            self._add_html_table(doc, element)
            
        elif tag == 'ul':
            self._add_list(doc, element, ordered=False)
            
        elif tag == 'ol':
            self._add_list(doc, element, ordered=True)
            
        elif tag == 'blockquote':
            para = doc.add_paragraph()
            # Apply blockquote style using mapper
            word_style = self.style_mapper.get_style_for_html_tag(tag)
            if word_style:
                self.style_mapper.apply_style_to_paragraph(para, word_style)
            else:
                para.style = 'Quote'  # Fallback
            self._process_inline_elements(para, element)
            
        elif tag == 'hr':
            # Add horizontal rule as empty paragraph
            para = doc.add_paragraph()
            para.add_run("─" * 50)
            
        # Handle text content after the element
        if element.tail and element.tail.strip():
            doc.add_paragraph(element.tail.strip())
            
        # Process child elements
        for child in element:
            self._process_html_element(doc, child, parent_para)
    
    def _process_inline_elements(self, para, element: ET.Element) -> None:
        """Process inline HTML elements within a paragraph."""
        if element.text:
            para.add_run(element.text)
            
        for child in element:
            tag = child.tag.lower()
            text = self._get_element_text(child)
            
            if tag == 'br':
                para.add_run().add_break()
            else:
                run = para.add_run(text)
                
                # Apply style using mapper
                word_style = self.style_mapper.get_style_for_html_tag(tag)
                if word_style:
                    self.style_mapper.apply_style_to_run(run, word_style)
                else:
                    # Fallback to original logic
                    if tag == 'strong' or tag == 'b':
                        run.bold = True
                    elif tag == 'em' or tag == 'i':
                        run.italic = True
                    elif tag == 'code':
                        run.font.name = 'Courier New'
                        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
                    elif tag == 'del' or tag == 's':
                        run.font.strike = True
                    elif tag == 'a':
                        run.font.color.rgb = RGBColor(0, 0, 255)
                        run.underline = True
                
            if child.tail:
                para.add_run(child.tail)
    
    def _add_code_block(self, doc: Document, element: ET.Element) -> None:
        """Add code block with proper formatting."""
        code_text = self._get_element_text(element)
        language = element.get('class', '').replace('language-', '') if element.get('class') else ''
        
        para = doc.add_paragraph(code_text)
        
        # Apply code block style using mapper
        word_style = self.style_mapper.get_word_style(MarkdownElement.CODE_BLOCK)
        if word_style:
            self.style_mapper.apply_style_to_paragraph(para, word_style)
            for run in para.runs:
                self.style_mapper.apply_style_to_run(run, word_style)
        else:
            # Fallback formatting
            for run in para.runs:
                run.font.name = 'Courier New'
                run.font.size = Inches(0.11)  # 10pt
            
        # Add language label if available
        if language:
            lang_para = doc.add_paragraph(f'Language: {language}')
            lang_para.style = 'Caption'
    
    def _add_html_table(self, doc: Document, table_element: ET.Element) -> None:
        """Add HTML table to document."""
        rows = []
        
        # Process table rows
        for row_elem in table_element.findall('.//tr'):
            row_data = []
            for cell_elem in row_elem.findall('.//td') + row_elem.findall('.//th'):
                cell_text = self._get_element_text(cell_elem)
                row_data.append(cell_text)
            if row_data:
                rows.append(row_data)
        
        if not rows:
            return
            
        # Create Word table
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Populate table
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(table.rows[row_idx].cells):
                    table.rows[row_idx].cells[col_idx].text = cell_data
                    
        # Style header row (first row is typically header)
        if rows:
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    def _add_list(self, doc: Document, list_element: ET.Element, ordered: bool = False) -> None:
        """Add ordered or unordered list."""
        for li_elem in list_element.findall('li'):
            text = self._get_element_text(li_elem)
            para = doc.add_paragraph(text)
            # Apply list style using mapper
            element = MarkdownElement.UNORDERED_LIST if not ordered else MarkdownElement.ORDERED_LIST
            word_style = self.style_mapper.get_word_style(element)
            if word_style:
                self.style_mapper.apply_style_to_paragraph(para, word_style)
            else:
                # Fallback
                para.style = 'List Bullet' if not ordered else 'List Number'
    
    def _get_element_text(self, element: ET.Element) -> str:
        """Get all text content from an element."""
        text_parts = []
        if element.text:
            text_parts.append(element.text)
        for child in element:
            text_parts.append(self._get_element_text(child))
            if child.tail:
                text_parts.append(child.tail)
        return ''.join(text_parts)
    
    def _fallback_text_conversion(self, doc: Document, html_content: str) -> None:
        """Fallback method to extract readable content from HTML while preserving basic formatting."""
        # Process line by line to preserve structure
        lines = html_content.split('\n')
        current_para = None
        
        for line in lines:
            line = line.strip()
            if not line:
                current_para = None
                continue
                
            # Handle headings
            heading_match = re.match(r'<h([1-6])[^>]*>(.*?)</h[1-6]>', line)
            if heading_match:
                level = int(heading_match.group(1))
                text = re.sub(r'<[^>]+>', '', heading_match.group(2))
                text = text.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
                doc.add_heading(text, level=level)
                current_para = None
                continue
            
            # Handle other content
            clean_text = re.sub(r'<[^>]+>', '', line)
            clean_text = clean_text.replace('&lt;', '<').replace('&gt;', '>').replace('&amp;', '&')
            
            if clean_text:
                if current_para is None:
                    current_para = doc.add_paragraph()
                if current_para.text:
                    current_para.add_run(' ')
                current_para.add_run(clean_text)
    
    def validate_markdown(self, markdown_content: str) -> Dict[str, Any]:
        """Validate markdown content for conversion compatibility."""
        try:
            # Pre-process content
            processed_content = self._preprocess_markdown(markdown_content)
            
            # Test parsing
            html_output = self.md.convert(processed_content)
            
            # Check for supported features
            features = {
                'tables': '<table>' in html_output,
                'code_blocks': '<pre>' in html_output or '<code>' in html_output,
                'headers': any(f'<h{i}>' in html_output for i in range(1, 7)),
                'lists': '<ul>' in html_output or '<ol>' in html_output,
                'emphasis': '<em>' in html_output or '<strong>' in html_output,
                'links': '<a href=' in html_output,
                'images': '<img' in html_output,
                'task_lists': '☑' in processed_content or '☐' in processed_content,
                'strikethrough': '<del>' in html_output,
            }
            
            return {
                'valid': True,
                'features_detected': features,
                'html_preview': html_output[:500] + '...' if len(html_output) > 500 else html_output
            }
            
        except Exception as e:
            return {
                'valid': False,
                'error': str(e),
                'features_detected': {},
                'html_preview': ''
            }