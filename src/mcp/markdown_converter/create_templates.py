"""Create basic Word templates for the converter."""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def create_simple_template():
    """Create a simple, clean template."""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add placeholder content
    doc.add_heading('Document Title', 0)
    doc.add_paragraph('This is a simple template for general documents.')
    
    return doc


def create_corporate_template():
    """Create a corporate-style template."""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.25)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Add corporate header
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_para.add_run('CORPORATE DOCUMENT')
    run.bold = True
    run.font.size = Inches(0.16)  # 12pt
    
    doc.add_paragraph()  # Spacing
    doc.add_heading('Document Title', 0)
    doc.add_paragraph('Professional corporate document template.')
    
    return doc


def create_technical_template():
    """Create a technical documentation template."""
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add technical header
    doc.add_heading('Technical Documentation', 0)
    doc.add_paragraph('Template for technical specifications and documentation.')
    
    # Add sample sections
    doc.add_heading('Overview', 1)
    doc.add_paragraph('Technical overview section.')
    
    doc.add_heading('Implementation', 1)
    doc.add_paragraph('Implementation details section.')
    
    return doc


def main():
    """Generate all template files."""
    template_dir = Path(__file__).parent / "templates"
    template_dir.mkdir(exist_ok=True)
    
    # Create templates
    templates = {
        'simple': create_simple_template(),
        'corporate': create_corporate_template(),
        'technical': create_technical_template()
    }
    
    # Save templates
    for name, doc in templates.items():
        path = template_dir / f"{name}.docx"
        doc.save(str(path))
        print(f"Created template: {path}")


if __name__ == "__main__":
    main()