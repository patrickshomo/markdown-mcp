"""Core markdown to DOCX conversion functionality."""

from typing import Optional, Dict, Any
from docx import Document
from pathlib import Path


class MarkdownConverter:
    """Converts markdown content to Word documents."""
    
    def __init__(self, template_path: Optional[str] = None):
        self.template_path = template_path
        
    def convert(self, markdown_content: str, metadata: Optional[Dict[str, Any]] = None) -> Document:
        """Convert markdown content to Word document."""
        if self.template_path and Path(self.template_path).exists():
            doc = Document(self.template_path)
        else:
            doc = Document()
            
        if metadata:
            self._add_metadata(doc, metadata)
            
        self._convert_content(doc, markdown_content)
        return doc
    
    def _add_metadata(self, doc: Document, metadata: Dict[str, Any]) -> None:
        """Add document metadata."""
        core_props = doc.core_properties
        if 'title' in metadata:
            core_props.title = metadata['title']
        if 'author' in metadata:
            core_props.author = metadata['author']
            
    def _convert_content(self, doc: Document, content: str) -> None:
        """Convert markdown content to Word paragraphs."""
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                doc.add_paragraph()
                i += 1
                continue
                
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('# ').strip()
                doc.add_heading(text, level=min(level, 9))
                i += 1
                continue
                
            if line.startswith('```'):
                i = self._add_code_block(doc, lines, i)
                continue
                
            doc.add_paragraph(line)
            i += 1
            
    def _add_code_block(self, doc: Document, lines: list, start_idx: int) -> int:
        """Add code block to document."""
        code_lines = []
        i = start_idx + 1
        
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
            
        code_text = '\n'.join(code_lines)
        para = doc.add_paragraph(code_text)
        # Use monospace font for code blocks
        for run in para.runs:
            run.font.name = 'Courier New'
        
        return i + 1 if i < len(lines) else i