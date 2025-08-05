"""Create DOTX templates from existing DOCX templates."""

from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_dotx_templates():
    """Create DOTX template files."""
    templates_dir = Path(__file__).parent / "templates"
    
    # Simple template
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Inches(0.11)
    doc.save(templates_dir / "simple.dotx")
    
    # Corporate template
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Inches(0.11)
    # Add corporate styling
    heading1 = doc.styles['Heading 1']
    heading1.font.name = 'Arial'
    heading1.font.size = Inches(0.16)
    doc.save(templates_dir / "corporate.dotx")
    
    # Technical template
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Inches(0.11)
    # Add technical styling
    heading1 = doc.styles['Heading 1']
    heading1.font.name = 'Times New Roman'
    heading1.font.size = Inches(0.14)
    doc.save(templates_dir / "technical.dotx")
    
    print("DOTX templates created successfully")


if __name__ == "__main__":
    create_dotx_templates()