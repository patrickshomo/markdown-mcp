"""Improved markdown to DOCX conversion using proper markdown libraries."""

from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX
from pathlib import Path
import markdown
from markdown.extensions import codehilite, tables, toc, fenced_code
# Import pymdown extensions properly
try:
    import pymdownx.tasklist
    import pymdownx.strikethrough
    import pymdownx.superfences
    import pymdownx.emoji
except ImportError:
    print("Warning: pymdownx extensions not available, some features may be limited")
import xml.etree.ElementTree as ET
from .template_manager import TemplateManager
from .image_handler import ImageHandler
from .mermaid_handler import MermaidHandler


class ImprovedMarkdownConverter:
    """Enhanced markdown to DOCX converter using proper markdown parsing."""
    
    def __init__(self, template_path: Optional[str] = None, template_dirs: Optional[list] = None, base_path: Optional[str] = None):
        self.template_path = template_path
        self.template_manager = TemplateManager(template_dirs)
        self.image_handler = ImageHandler(base_path)
        self.mermaid_handler = MermaidHandler()
        
        # Configure markdown parser with GFM extensions
        self.md = markdown.Markdown(extensions=[
            'tables',           # GFM tables
            'fenced_code',      # Fenced code blocks
            'codehilite',       # Code highlighting
            'toc',              # Table of contents
            'nl2br',            # Newline to break
            'pymdownx.tasklist',     # Task lists
            'pymdownx.strikethrough', # Strikethrough
            'pymdownx.superfences',   # Enhanced fenced code
            'pymdownx.emoji',         # Emoji support
        ], extension_configs={
            'pymdownx.tasklist': {
                'custom_checkbox': True,
                'clickable_checkbox': False,
            },
            'pymdownx.superfences': {
                'custom_fences': [
                    {
                        'name': 'mermaid',
                        'class': 'mermaid',
                        'format': self._format_mermaid
                    }
                ]
            },
            'codehilite': {
                'use_pygments': False,
                'noclasses': True,
            }
        })
        
    def _format_mermaid(self, source, language, css_class, options, md, **kwargs):
        """Custom formatter for mermaid diagrams."""
        return f'<div class="mermaid">{source}</div>'
        
    def convert(self, markdown_content: str, metadata: Optional[Dict[str, Any]] = None, template_name: Optional[str] = None) -> Document:
        """Convert markdown content to Word document using proper parsing."""
        template_path = None
        if template_name:
            template_path = self.template_manager.get_template_path(template_name)
        elif self.template_path:
            template_path = self.template_path
            
        if template_path and self.template_manager.validate_template(template_path):
            doc = Document(template_path)
        else:
            doc = Document()
            
        if metadata:
            self._add_metadata(doc, metadata)
            
        # Process images and mermaid diagrams first
        processed_content = self.image_handler.process_images(doc, markdown_content)
        processed_content = self.mermaid_handler.process_mermaid(doc, processed_content)
        
        # Parse markdown to HTML
        html_content = self.md.convert(processed_content)
        
        # Convert HTML to Word document
        self._convert_html_to_docx(doc, html_content)
        
        return doc
    
    def _add_metadata(self, doc: Document, metadata: Dict[str, Any]) -> None:
        """Add document metadata."""
        core_props = doc.core_properties
        if 'title' in metadata:
            core_props.title = metadata['title']
        if 'author' in metadata:
            core_props.author = metadata['author']
        if 'subject' in metadata:
            core_props.subject = metadata['subject']
        if 'keywords' in metadata:
            core_props.keywords = metadata['keywords']
        if 'comments' in metadata:
            core_props.comments = metadata['comments']
    
    def _convert_html_to_docx(self, doc: Document, html_content: str) -> None:
        """Convert HTML to Word document elements."""
        try:
            # Parse HTML
            root = ET.fromstring(f'<root>{html_content}</root>')
            self._process_html_element(doc, root)
        except ET.ParseError:
            # Fallback to simple text if HTML parsing fails
            doc.add_paragraph(html_content)
    
    def _process_html_element(self, doc: Document, element: ET.Element, parent_para=None) -> None:
        """Process HTML elements recursively."""
        tag = element.tag.lower()
        
        if tag in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(tag[1])
            text = self._get_element_text(element)
            doc.add_heading(text, level=level)
            
        elif tag == 'p':
            para = doc.add_paragraph()
            self._process_inline_elements(para, element)
            
        elif tag == 'pre':
            code_element = element.find('code')
            if code_element is not None:
                self._add_code_block(doc, code_element)
            else:
                self._add_code_block(doc, element)
                
        elif tag == 'table':
            self._add_html_table(doc, element)
            
        elif tag == 'ul':
            self._add_list(doc, element, ordered=False)
            
        elif tag == 'ol':
            self._add_list(doc, element, ordered=True)
            
        elif tag == 'blockquote':
            para = doc.add_paragraph()
            para.style = 'Quote'
            self._process_inline_elements(para, element)
            
        elif tag == 'hr':
            # Add horizontal rule as empty paragraph with border
            para = doc.add_paragraph()
            para.paragraph_format.border_bottom.color.rgb = RGBColor(0, 0, 0)
            
        elif tag == 'div' and element.get('class') == 'mermaid':
            # Handle mermaid diagrams
            mermaid_code = self._get_element_text(element)
            self.mermaid_handler.process_mermaid(doc, f'```mermaid\\n{mermaid_code}\\n```')
            
        # Process child elements
        for child in element:
            self._process_html_element(doc, child, parent_para)
    
    def _process_inline_elements(self, para, element: ET.Element) -> None:
        """Process inline HTML elements within a paragraph."""
        if element.text:
            para.add_run(element.text)
            
        for child in element:
            tag = child.tag.lower()
            text = self._get_element_text(child)
            
            if tag == 'strong' or tag == 'b':
                run = para.add_run(text)
                run.bold = True
            elif tag == 'em' or tag == 'i':
                run = para.add_run(text)
                run.italic = True
            elif tag == 'code':
                run = para.add_run(text)
                run.font.name = 'Courier New'
                run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
            elif tag == 'del' or tag == 's':
                run = para.add_run(text)
                run.font.strike = True
            elif tag == 'a':
                # Handle links
                href = child.get('href', '')
                run = para.add_run(text)
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
                # Note: Word hyperlinks require more complex handling
            elif tag == 'input' and child.get('type') == 'checkbox':
                # Handle task list checkboxes
                checked = child.get('checked') is not None
                checkbox_text = '☑' if checked else '☐'
                para.add_run(f'{checkbox_text} ')
            else:
                para.add_run(text)
                
            if child.tail:
                para.add_run(child.tail)
    
    def _add_code_block(self, doc: Document, element: ET.Element) -> None:
        """Add code block with proper formatting."""
        code_text = self._get_element_text(element)
        language = element.get('class', '').replace('language-', '') if element.get('class') else ''
        
        para = doc.add_paragraph(code_text)
        
        # Enhanced code block formatting
        for run in para.runs:
            run.font.name = 'Courier New'
            run.font.size = Inches(0.11)  # 10pt
            
        # Add language label if available
        if language:
            lang_para = doc.add_paragraph(f'Language: {language}')
            lang_para.style = 'Caption'
    
    def _add_html_table(self, doc: Document, table_element: ET.Element) -> None:
        """Add HTML table to document."""
        rows = []
        
        # Process table rows
        for row_elem in table_element.findall('.//tr'):
            row_data = []
            for cell_elem in row_elem.findall('.//td') + row_elem.findall('.//th'):
                cell_text = self._get_element_text(cell_elem)
                row_data.append(cell_text)
            if row_data:
                rows.append(row_data)
        
        if not rows:
            return
            
        # Create Word table
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Populate table
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(table.rows[row_idx].cells):
                    table.rows[row_idx].cells[col_idx].text = cell_data
                    
        # Style header row (first row is typically header)
        if rows:
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
    
    def _add_list(self, doc: Document, list_element: ET.Element, ordered: bool = False) -> None:
        """Add ordered or unordered list."""
        for li_elem in list_element.findall('li'):
            text = self._get_element_text(li_elem)
            
            # Check for task list items
            if li_elem.find('.//input[@type="checkbox"]') is not None:
                checkbox = li_elem.find('.//input[@type="checkbox"]')
                checked = checkbox.get('checked') is not None
                checkbox_text = '☑' if checked else '☐'
                para = doc.add_paragraph(f'{checkbox_text} {text}')
            else:
                para = doc.add_paragraph(text, style='List Bullet' if not ordered else 'List Number')
    
    def _get_element_text(self, element: ET.Element) -> str:
        """Get all text content from an element."""
        text_parts = []
        if element.text:
            text_parts.append(element.text)
        for child in element:
            text_parts.append(self._get_element_text(child))
            if child.tail:
                text_parts.append(child.tail)
        return ''.join(text_parts)
    
    def validate_markdown(self, markdown_content: str) -> Dict[str, Any]:
        """Validate markdown content for conversion compatibility."""
        try:
            # Test parsing
            html_output = self.md.convert(markdown_content)
            
            # Check for supported features
            features = {
                'tables': '<table>' in html_output,
                'code_blocks': '<pre>' in html_output or '<code>' in html_output,
                'headers': any(f'<h{i}>' in html_output for i in range(1, 7)),
                'lists': '<ul>' in html_output or '<ol>' in html_output,
                'emphasis': '<em>' in html_output or '<strong>' in html_output,
                'links': '<a href=' in html_output,
                'images': '<img' in html_output,
                'task_lists': 'type="checkbox"' in html_output,
                'strikethrough': '<del>' in html_output or '<s>' in html_output,
            }
            
            return {
                'valid': True,
                'features_detected': features,
                'html_preview': html_output[:500] + '...' if len(html_output) > 500 else html_output
            }
            
        except Exception as e:
            return {
                'valid': False,
                'error': str(e),
                'features_detected': {},
                'html_preview': ''
            }