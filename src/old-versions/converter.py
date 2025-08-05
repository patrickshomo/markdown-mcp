"""Core markdown to DOCX conversion functionality."""

from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
import re
from .template_manager import TemplateManager
from .image_handler import ImageHandler
from .mermaid_handler import MermaidHandler


class MarkdownConverter:
    """Converts markdown content to Word documents."""
    
    def __init__(self, template_path: Optional[str] = None, template_dirs: Optional[list] = None, base_path: Optional[str] = None):
        self.template_path = template_path
        self.template_manager = TemplateManager(template_dirs)
        self.image_handler = ImageHandler(base_path)
        self.mermaid_handler = MermaidHandler()
        
    def convert(self, markdown_content: str, metadata: Optional[Dict[str, Any]] = None, template_name: Optional[str] = None) -> Document:
        """Convert markdown content to Word document."""
        template_path = None
        if template_name:
            template_path = self.template_manager.get_template_path(template_name)
        elif self.template_path:
            template_path = self.template_path
            
        if template_path and self.template_manager.validate_template(template_path):
            doc = Document(template_path)
        else:
            doc = Document()
            
        if metadata:
            self._add_metadata(doc, metadata)
            
        # Process images and mermaid diagrams
        processed_content = self.image_handler.process_images(doc, markdown_content)
        processed_content = self.mermaid_handler.process_mermaid(doc, processed_content)
            
        self._convert_content(doc, processed_content)
        return doc
    
    def _add_metadata(self, doc: Document, metadata: Dict[str, Any]) -> None:
        """Add document metadata."""
        core_props = doc.core_properties
        if 'title' in metadata:
            core_props.title = metadata['title']
        if 'author' in metadata:
            core_props.author = metadata['author']
        if 'subject' in metadata:
            core_props.subject = metadata['subject']
        if 'keywords' in metadata:
            core_props.keywords = metadata['keywords']
        if 'comments' in metadata:
            core_props.comments = metadata['comments']
            
    def _convert_content(self, doc: Document, content: str) -> None:
        """Convert markdown content to Word paragraphs."""
        lines = content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                doc.add_paragraph()
                i += 1
                continue
                
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('# ').strip()
                doc.add_heading(text, level=min(level, 9))
                i += 1
                continue
                
            if line.startswith('```'):
                i = self._add_code_block(doc, lines, i)
                continue
                
            if self._is_table_row(line):
                i = self._add_table(doc, lines, i)
                continue
                
            self._add_formatted_paragraph(doc, line)
            i += 1
            
    def _add_code_block(self, doc: Document, lines: list, start_idx: int) -> int:
        """Add code block to document."""
        code_lines = []
        i = start_idx + 1
        
        while i < len(lines) and not lines[i].strip().startswith('```'):
            code_lines.append(lines[i])
            i += 1
            
        code_text = '\n'.join(code_lines)
        para = doc.add_paragraph(code_text)
        
        # Enhanced code block formatting
        for run in para.runs:
            run.font.name = 'Courier New'
            run.font.size = Inches(0.11)  # 10pt
        
        # Add subtle background styling
        para.style = 'No Spacing'
        
        return i + 1 if i < len(lines) else i
    
    def _is_table_row(self, line: str) -> bool:
        """Check if line is a markdown table row."""
        return '|' in line and line.strip().startswith('|') and line.strip().endswith('|')
    
    def _add_table(self, doc: Document, lines: list, start_idx: int) -> int:
        """Add markdown table to document."""
        table_lines = []
        i = start_idx
        
        # Collect all table rows
        while i < len(lines) and self._is_table_row(lines[i]):
            table_lines.append(lines[i].strip())
            i += 1
            
        if len(table_lines) < 2:
            return i
            
        # Parse table data
        rows = []
        for line in table_lines:
            # Skip separator rows (contains only |, -, :, spaces)
            if re.match(r'^\|[\s\-:]+\|$', line):
                continue
            cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if cells:
                rows.append(cells)
                
        if not rows:
            return i
            
        # Create Word table
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        
        # Populate table
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(table.rows[row_idx].cells):
                    table.rows[row_idx].cells[col_idx].text = cell_data
                    
        # Style header row
        if rows:
            for cell in table.rows[0].cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        
        return i
    
    def _add_formatted_paragraph(self, doc: Document, text: str) -> None:
        """Add paragraph with inline markdown formatting."""
        para = doc.add_paragraph()
        self._process_inline_formatting(para, text)
    
    def _process_inline_formatting(self, para, text: str) -> None:
        """Process inline markdown formatting in text."""
        # Pattern for markdown formatting: **bold**, *italic*, `code`, ~~strike~~
        pattern = r'(\*\*.*?\*\*)|(\*.*?\*)|(~~.*?~~)|(`.+?`)|([^*~`]+)'
        
        for match in re.finditer(pattern, text):
            content = match.group(0)
            
            if content.startswith('**') and content.endswith('**'):
                # Bold text
                run = para.add_run(content[2:-2])
                run.bold = True
            elif content.startswith('*') and content.endswith('*'):
                # Italic text
                run = para.add_run(content[1:-1])
                run.italic = True
            elif content.startswith('~~') and content.endswith('~~'):
                # Strikethrough text
                run = para.add_run(content[2:-2])
                run.font.strike = True
            elif content.startswith('`') and content.endswith('`'):
                # Inline code
                run = para.add_run(content[1:-1])
                run.font.name = 'Courier New'
            else:
                # Regular text
                para.add_run(content)