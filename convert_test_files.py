#!/usr/bin/env python3
"""Convert test input files to output."""

import os
import sys
from pathlib import Path

# Add src to path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), 'src'))

try:
    from mcp.markdown_converter.simple_improved_converter import SimpleImprovedConverter
except ImportError:
    # Fallback - create minimal converter
    import markdown
    from docx import Document
    
    class SimpleImprovedConverter:
        def __init__(self):
            self.md = markdown.Markdown(extensions=['tables', 'fenced_code'])
        
        def convert(self, content, metadata=None):
            doc = Document()
            if metadata and 'title' in metadata:
                doc.core_properties.title = metadata['title']
            
            html = self.md.convert(content)
            # Simple conversion - just add as paragraphs
            lines = content.split('\n')
            for line in lines:
                if line.strip():
                    if line.startswith('#'):
                        level = len(line) - len(line.lstrip('#'))
                        text = line.lstrip('# ').strip()
                        doc.add_heading(text, level=min(level, 6))
                    else:
                        doc.add_paragraph(line)
            return doc

def convert_files():
    input_dir = Path("tests/input")
    output_dir = Path("tests/output")
    output_dir.mkdir(exist_ok=True)
    
    converter = SimpleImprovedConverter()
    
    for md_file in input_dir.glob("*.md"):
        print(f"Converting {md_file.name}...")
        
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        doc = converter.convert(content, {"title": md_file.stem})
        output_file = output_dir / f"{md_file.stem}.docx"
        doc.save(str(output_file))
        
        print(f"Saved to {output_file}")

if __name__ == "__main__":
    convert_files()